import re
from progress.bar import IncrementalBar
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Функция обхода классных показателей в скобках
def brackets_clean(string):
	pos_ns = string.find('(')
	pos_ks = string.find(')')
	kl_pk = string[pos_ns:pos_ks+1]
	string = re.sub(' \(.*\)', '', string)
	word_list = string.split(',')
	word_list[0] = word_list[0] + ' ' + kl_pk

	return word_list

# Функция добавления части речи (гл.)
def add_part_speech(text):
	pr_list = text.split('\n')

	for pr_list_item in pr_list:
		current_pr = pr_list.index(pr_list_item)
		word_list = pr_list_item.split(',')

		if re.search(' \([в,б,д,й],.*\)', pr_list_item):
			word_list = brackets_clean(pr_list_item)

		if len(word_list) > 139:
			word_list[0] = word_list[0] + ' (гл.)'
			pr_list[current_pr] = ','.join(word_list)

	text_part_speech = '\n'.join(pr_list)

	return text_part_speech


def homonyms(text):

	homonyms_list = {}


	# Очищаем строку от пробелов в начале и конце
	text = text.strip()

	# Добавляем в нужных местах часть речи (гл.)
	text = add_part_speech(text)

	# Преобразуем текст в список по абзацам
	pr_list = text.split('\n')
	
	# Индикатор выполнения
	bar = IncrementalBar('homonyms', max = len(pr_list))

	# Перебираем каждый абзац текста
	for pr_list_item in pr_list:

		bar.next()

		# Индекс текущего абзаца
		current_pr = pr_list.index(pr_list_item)

		# Преобразуем каждый абзац текста в список по словам
		word_list = pr_list_item.split(',')

		# Если есть классный показатель в скобках,
		# убираем его, делим строку в список по запятым,
		# и обратно добавляем к 1-му элементу списка классный показатель
		if re.search(' \([в,б,д,й],.*\)', pr_list_item):
			word_list = brackets_clean(pr_list_item)

		# Переводим все буквы слов в нижний регистр
		pr_list_item_lower = pr_list_item.lower()

		#Если есть цифры, удаляем 
		if re.search('\d+', pr_list_item_lower):
			pr_list_item_lower = re.sub('\d+', '', pr_list_item_lower)

		#Если есть скобки, удаляем
		if re.search(' \(.*\)', pr_list_item_lower):
			pr_list_item_lower = re.sub(' \(.*\)', '', pr_list_item_lower)

		# Преобразуем каждый абзац текста в список по словам в нижнем регистре
		word_list_lower = pr_list_item_lower.split(',')


		# Еще раз перебираем каждый абзац
		# для сравнения двух абзацев
		for pr_list_item_2 in pr_list:

			# Индекс текущего абзаца
			current_pr_2 = pr_list.index(pr_list_item_2)
		

			if current_pr != current_pr_2:

				# Преобразуем каждый абзац текста в список по словам
				word_list_2 = pr_list_item_2.split(',')

				# Если есть классный показатель в скобках,
				# убираем его, делим строку в список по запятым,
				# и обратно добавляем к 1 элементу списка классный показатель
				if re.search(' \([в,б,д,й],.*\)', pr_list_item_2):
					word_list_2 = brackets_clean(pr_list_item_2)

				# Переводим все буквы слов в нижний регистр
				pr_list_item_2_lower = pr_list_item_2.lower()

				#Если есть цифры, удаляем 
				if re.search('\d+', pr_list_item_2_lower):
					pr_list_item_2_lower = re.sub('\d+', '', pr_list_item_2_lower)

				#Если есть скобки, удаляем
				if re.search(' \(.*\)', pr_list_item_2_lower):
					pr_list_item_2_lower = re.sub(' \(.*\)', '', pr_list_item_2_lower)

				# Преобразуем каждый абзац текста в список по словам в нижнем регистре
				word_list_2_lower = pr_list_item_2_lower.split(',')

				# Получаем омонимы сравнивая два абзаца
				om_temp_list = list(set(word_list_lower) & set(word_list_2_lower))
				
				if '#' in word_list_lower[0] and '#' in word_list_2_lower[0]:
					om_temp_list = list(set(word_list_lower[1:]) & set(word_list_2_lower[1:]))

				if om_temp_list:
					for homonym in om_temp_list:

						# Получаем индекс главного слова-омонима
						word_index = word_list_lower.index(homonym)

						# Получаем главное слово-омоним по индексу
						word = word_list[word_index]

						# Получаем индекс слова-омонима
						word_homonym_index = word_list_2_lower.index(homonym)

						# Получаем само слово-омоним по индексу
						word_homonym = word_list_2[word_homonym_index]

						# Проверяем является ли слово главным словом абзаца,
						# если не является, добавляем 'от "главное слово абзаца"' 
						if word_index > 0:
							word_glav = word + ' от ' + word_list[0]
								
							# Удаляем символ "#", если есть
							word_glav = re.sub('#', '', word_glav)
						else:
							word_glav = word + ' основное слово'

						# Проверяем является ли найденный омоним главным словом абзаца,
						# если не является, добавляем 'от "главное слово абзаца"' 
						if word_homonym_index > 0:
							word_homonym = word_homonym + ' от ' + word_list_2[0]

							# Удаляем символ "#", если есть
							word_homonym = re.sub('#', '', word_homonym)
						else:
							word_homonym = word_homonym + ' основное слово'

						
							
						# Проверяем существование главного слова в словаре омонимов
						if homonym in homonyms_list:
							# Если есть, добавляем омоним этого слова
							homonyms_list[homonym].append(word_homonym)
						else:
							# Если нет, добавляем само слово и его омоним в словарь
							homonyms_list.update({homonym:[word_glav]})
							homonyms_list[homonym].append(word_homonym)

							
						# Удаляем из основного списка найденный омоним, если слово не является "Главным"
						if word_homonym_index == 0:
							word_list_2[word_homonym_index] = word_list_2[word_homonym_index] + '#'
							word_list_2_lower[word_homonym_index] = word_list_2_lower[word_homonym_index] + '#'
							pr_list[current_pr_2] = ','.join(word_list_2)

						while homonym in word_list_2_lower:
							index_for_delete = word_list_2_lower.index(homonym)
							try:
								word_list_2.pop(index_for_delete)
								word_list_2_lower.pop(index_for_delete)
							except IndexError:
								pass
							else:
								pr_list[current_pr_2] = ','.join(word_list_2)
				
	bar.finish()

	return homonyms_list

if __name__ == '__main__':

	begin = time.time()

	with open("dictionary.txt", 'r', encoding="utf-8") as file:
	    dictionary = file.read()
	
	homonyms = homonyms(dictionary)

	# Запись результата в документ Word
	homonyms_doc = Document()
	for key in homonyms:
	    homonym_glav = key 
	    homonym_2_list = homonyms[key]  
	    p = homonyms_doc.add_paragraph()
	    p_format = p.paragraph_format
	    p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	    run = p.add_run(homonym_glav.upper())
	    run.bold = True
	    run.font.size = Pt(16)

	    for homonym_2 in homonym_2_list:
	    	p_2 = homonyms_doc.add_paragraph()
	    	if ' от ' in homonym_2:
	    		pos_ot = homonym_2.find(' от ')
	    		homonym_start = homonym_2[:pos_ot]
	    		homonym_end = homonym_2[pos_ot:]
	    		run_om_1 = p_2.add_run(homonym_start)
	    		run_om_1.bold = True
	    		run_om_1.font.size = Pt(14)
	    		run_om_2 = p_2.add_run(homonym_end)
	    		run_om_2.font.size = Pt(14)
	    	elif ' основное слово' in homonym_2:
	    		pos_os = homonym_2.find(' основное слово')
	    		homonym_start = homonym_2[:pos_os]
	    		homonym_end = homonym_2[pos_os:]
	    		run_om_1 = p_2.add_run(homonym_start)
	    		run_om_1.bold = True
	    		run_om_1.font.size = Pt(14)
	    		run_om_2 = p_2.add_run(homonym_end)
	    		run_om_2.font.size = Pt(14)
	    	else:
	    		run_om_1 = p_2.add_run(homonym_2)
	    		run_om_1.bold = True
	    		run_om_1.font.size = Pt(14)
	    		
	    homonyms_doc.add_paragraph()	
	homonyms_doc.save('homonyms.docx')
    
	# Расчет времени
	end = time.time()
	t = end-begin
	str_t = str(t) + ' c.'
	if t > 60 and t < 3600:
		t = t/60
		str_t = str(t) + ' м.'
	elif t > 3600:
		t = t/60/60
		str_t = str(t) + ' ч.'
	print('Время: ' + str_t)
	print('Готово')

